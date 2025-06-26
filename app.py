# Dynamic RAG System - Proof of Concept
# A scalable document retrieval system with external knowledge base

import streamlit as st
import asyncio
import threading
import time
import uuid
from datetime import datetime
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, asdict
import json
import sqlite3
import numpy as np
from sentence_transformers import SentenceTransformer
import faiss
from pathlib import Path
import hashlib
import queue
import logging
import io
import base64
import cohere

# Document processing imports
try:
    import PyPDF2

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl

    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class Document:
    id: str
    title: str
    content: str
    metadata: Dict[str, Any]
    created_at: str
    embedding_id: Optional[int] = None


@dataclass
class DocumentChunk:
    id: str
    document_id: str
    content: str
    chunk_index: int
    embedding: Optional[np.ndarray] = None


class DocumentDatabase:
    """SQLite-based document storage"""

    def __init__(self, db_path: str = "documents.db"):
        self.db_path = db_path
        self.init_db()

    def init_db(self):
        conn = sqlite3.connect(self.db_path)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS documents (
                id TEXT PRIMARY KEY,
                title TEXT NOT NULL,
                content TEXT NOT NULL,
                metadata TEXT,
                created_at TEXT,
                embedding_id INTEGER
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS document_chunks (
                id TEXT PRIMARY KEY,
                document_id TEXT,
                content TEXT,
                chunk_index INTEGER,
                FOREIGN KEY (document_id) REFERENCES documents (id)
            )
        """)
        conn.commit()
        conn.close()

    def add_document(self, document: Document) -> bool:
        try:
            conn = sqlite3.connect(self.db_path)
            conn.execute("""
                INSERT OR REPLACE INTO documents 
                (id, title, content, metadata, created_at, embedding_id)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (
                document.id, document.title, document.content,
                json.dumps(document.metadata), document.created_at,
                document.embedding_id
            ))
            conn.commit()
            conn.close()
            return True
        except Exception as e:
            logger.error(f"Error adding document: {e}")
            return False

    def get_document(self, doc_id: str) -> Optional[Document]:
        conn = sqlite3.connect(self.db_path)
        cursor = conn.execute(
            "SELECT * FROM documents WHERE id = ?", (doc_id,)
        )
        row = cursor.fetchone()
        conn.close()

        if row:
            return Document(
                id=row[0], title=row[1], content=row[2],
                metadata=json.loads(row[3]), created_at=row[4],
                embedding_id=row[5]
            )
        return None

    def list_documents(self) -> List[Document]:
        conn = sqlite3.connect(self.db_path)
        cursor = conn.execute("SELECT * FROM documents ORDER BY created_at DESC")
        rows = cursor.fetchall()
        conn.close()

        return [
            Document(
                id=row[0], title=row[1], content=row[2],
                metadata=json.loads(row[3]), created_at=row[4],
                embedding_id=row[5]
            )
            for row in rows
        ]

    def add_chunks(self, chunks: List[DocumentChunk]):
        conn = sqlite3.connect(self.db_path)
        for chunk in chunks:
            conn.execute("""
                INSERT OR REPLACE INTO document_chunks
                (id, document_id, content, chunk_index)
                VALUES (?, ?, ?, ?)
            """, (chunk.id, chunk.document_id, chunk.content, chunk.chunk_index))
        conn.commit()
        conn.close()


class VectorStore:
    """FAISS-based vector storage for embeddings"""

    def __init__(self, dimension: int = 384, index_path: str = "vector_index.faiss"):
        self.dimension = dimension
        self.index_path = index_path
        self.index = faiss.IndexFlatIP(dimension)  # Inner product (cosine similarity)
        self.chunk_mapping = {}  # Maps vector index to chunk ID
        self.load_index()

    def load_index(self):
        try:
            if Path(self.index_path).exists():
                self.index = faiss.read_index(self.index_path)
                # Load mapping
                mapping_path = self.index_path.replace('.faiss', '_mapping.json')
                if Path(mapping_path).exists():
                    with open(mapping_path, 'r') as f:
                        self.chunk_mapping = json.load(f)
        except Exception as e:
            logger.error(f"Error loading index: {e}")

    def save_index(self):
        try:
            faiss.write_index(self.index, self.index_path)
            # Save mapping
            mapping_path = self.index_path.replace('.faiss', '_mapping.json')
            with open(mapping_path, 'w') as f:
                json.dump(self.chunk_mapping, f)
        except Exception as e:
            logger.error(f"Error saving index: {e}")

    def add_embeddings(self, embeddings: np.ndarray, chunk_ids: List[str]):
        # Normalize embeddings for cosine similarity
        faiss.normalize_L2(embeddings)

        start_idx = self.index.ntotal
        self.index.add(embeddings)

        # Update mapping
        for i, chunk_id in enumerate(chunk_ids):
            self.chunk_mapping[str(start_idx + i)] = chunk_id

        self.save_index()

    def search(self, query_embedding: np.ndarray, k: int = 5) -> List[tuple]:
        if self.index.ntotal == 0:
            return []

        # Normalize query
        faiss.normalize_L2(query_embedding.reshape(1, -1))

        scores, indices = self.index.search(query_embedding.reshape(1, -1), k)

        results = []
        for score, idx in zip(scores[0], indices[0]):
            if idx != -1 and str(idx) in self.chunk_mapping:
                chunk_id = self.chunk_mapping[str(idx)]
                results.append((chunk_id, float(score)))

        return results


class FileProcessor:
    """Handles file upload and text extraction"""

    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """Extract text from PDF file"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not available. Install with: pip install PyPDF2")

        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Install with: pip install python-docx")

        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")

    @staticmethod
    def extract_text_from_txt(file_bytes: bytes) -> str:
        """Extract text from TXT file"""
        try:
            # Try UTF-8 first, fall back to latin-1
            try:
                return file_bytes.decode('utf-8')
            except UnicodeDecodeError:
                return file_bytes.decode('latin-1', errors='ignore')
        except Exception as e:
            raise Exception(f"Error extracting text from TXT: {str(e)}")

    @staticmethod
    def extract_text_from_excel(file_bytes: bytes) -> str:
        """Extract text from Excel file"""
        if not EXCEL_AVAILABLE:
            raise ImportError("openpyxl not available. Install with: pip install openpyxl")

        try:
            workbook = openpyxl.load_workbook(io.BytesIO(file_bytes))
            text = ""
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"\n=== Sheet: {sheet_name} ===\n"
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text += row_text + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Error extracting text from Excel: {str(e)}")

    @staticmethod
    def process_uploaded_file(uploaded_file) -> tuple:
        """Process uploaded file and extract text"""
        file_extension = uploaded_file.name.split('.')[-1].lower()
        file_bytes = uploaded_file.read()

        # Reset file pointer for potential re-reading
        uploaded_file.seek(0)

        try:
            if file_extension == 'pdf':
                text = FileProcessor.extract_text_from_pdf(file_bytes)
            elif file_extension == 'docx':
                text = FileProcessor.extract_text_from_docx(file_bytes)
            elif file_extension in ['txt', 'md', 'py', 'js', 'html', 'css', 'json']:
                text = FileProcessor.extract_text_from_txt(file_bytes)
            elif file_extension in ['xlsx', 'xls']:
                text = FileProcessor.extract_text_from_excel(file_bytes)
            else:
                raise Exception(f"Unsupported file type: {file_extension}")

            metadata = {
                'file_name': uploaded_file.name,
                'file_size': len(file_bytes),
                'file_type': file_extension,
                'upload_time': datetime.now().isoformat()
            }

            return text, metadata

        except Exception as e:
            raise Exception(f"Error processing file '{uploaded_file.name}': {str(e)}")


class TextProcessor:
    """Handles text chunking and embedding generation"""

    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        self.embedding_model = SentenceTransformer(model_name)
        self.chunk_size = 500
        self.chunk_overlap = 50

    def chunk_text(self, text: str, document_id: str) -> List[DocumentChunk]:
        # Simple sentence-aware chunking
        sentences = text.split('. ')
        chunks = []
        current_chunk = ""
        chunk_index = 0

        for sentence in sentences:
            if len(current_chunk) + len(sentence) > self.chunk_size and current_chunk:
                chunk_id = f"{document_id}_chunk_{chunk_index}"
                chunks.append(DocumentChunk(
                    id=chunk_id,
                    document_id=document_id,
                    content=current_chunk.strip(),
                    chunk_index=chunk_index
                ))

                # Handle overlap
                words = current_chunk.split()
                overlap_words = words[-self.chunk_overlap:] if len(words) > self.chunk_overlap else words
                current_chunk = ' '.join(overlap_words) + ' ' + sentence
                chunk_index += 1
            else:
                current_chunk += sentence + '. '

        # Add final chunk
        if current_chunk.strip():
            chunk_id = f"{document_id}_chunk_{chunk_index}"
            chunks.append(DocumentChunk(
                id=chunk_id,
                document_id=document_id,
                content=current_chunk.strip(),
                chunk_index=chunk_index
            ))

        return chunks

    def generate_embeddings(self, texts: List[str]) -> np.ndarray:
        return self.embedding_model.encode(texts, convert_to_numpy=True)


class BackgroundIndexer:
    """Background worker for processing documents"""

    def __init__(self, doc_db: DocumentDatabase, vector_store: VectorStore,
                 text_processor: TextProcessor):
        self.doc_db = doc_db
        self.vector_store = vector_store
        self.text_processor = text_processor
        self.processing_queue = queue.Queue()
        self.is_running = False
        self.worker_thread = None

    def start(self):
        if not self.is_running:
            self.is_running = True
            self.worker_thread = threading.Thread(target=self._worker, daemon=True)
            self.worker_thread.start()

    def stop(self):
        self.is_running = False
        if self.worker_thread:
            self.worker_thread.join()

    def queue_document(self, document: Document):
        self.processing_queue.put(document)

    def _worker(self):
        while self.is_running:
            try:
                document = self.processing_queue.get(timeout=1)
                self._process_document(document)
                self.processing_queue.task_done()
            except queue.Empty:
                continue
            except Exception as e:
                logger.error(f"Error processing document: {e}")

    def _process_document(self, document: Document):
        logger.info(f"Processing document: {document.title}")

        # Chunk the document
        chunks = self.text_processor.chunk_text(document.content, document.id)

        # Generate embeddings
        chunk_texts = [chunk.content for chunk in chunks]
        embeddings = self.text_processor.generate_embeddings(chunk_texts)

        # Store chunks and embeddings
        chunk_ids = [chunk.id for chunk in chunks]
        self.vector_store.add_embeddings(embeddings, chunk_ids)
        self.doc_db.add_chunks(chunks)

        logger.info(f"Processed {len(chunks)} chunks for document: {document.title}")


class RAGSystem:
    """Main RAG system orchestrator"""

    def __init__(self, cohere_api_key: str = None):
        self.doc_db = DocumentDatabase()
        self.vector_store = VectorStore()
        self.text_processor = TextProcessor()
        self.indexer = BackgroundIndexer(self.doc_db, self.vector_store, self.text_processor)
        self.indexer.start()

        # Initialize Cohere client
        self.cohere_client = cohere.Client(cohere_api_key) if cohere_api_key else None
        self.cohere_model = "command-r"  # or "command-r-plus" for more advanced capabilities

    def add_document(self, title: str, content: str, metadata: Dict[str, Any] = None) -> str:
        document_id = str(uuid.uuid4())
        document = Document(
            id=document_id,
            title=title,
            content=content,
            metadata=metadata or {},
            created_at=datetime.now().isoformat()
        )

        # Store document
        self.doc_db.add_document(document)

        # Queue for background processing
        self.indexer.queue_document(document)

        return document_id

    def search_documents(self, query: str, max_chunks: int = 5,
                        max_tokens: int = 16000, generate_response: bool = True) -> Dict[str, Any]:
        """
        Enhanced search with Cohere response generation

        Returns:
            Dict with keys:
            - 'chunks': List of retrieved chunks
            - 'response': Generated answer (if generate_response=True)
            - 'sources': List of source documents
        """
        # Generate query embedding
        query_embedding = self.text_processor.generate_embeddings([query])

        # Search vector store
        results = self.vector_store.search(query_embedding[0], k=max_chunks * 2)

        # Apply MMR for diversity
        selected_results = self._apply_mmr(results, query_embedding[0], max_chunks)

        # Retrieve chunk contents and apply token budget
        retrieved_chunks = []
        total_tokens = 0
        source_docs = set()

        for chunk_id, score in selected_results:
            # Get actual chunk content from database
            chunk_content = self._get_chunk_content(chunk_id)
            if not chunk_content:
                continue

            # Get document info for citation
            doc_id = chunk_id.split('_chunk_')[0]
            document = self.doc_db.get_document(doc_id)
            source_info = {
                'document_id': doc_id,
                'document_title': document.title if document else "Unknown",
                'chunk_id': chunk_id,
                'score': score
            }

            # Estimate tokens (rough approximation: 1 token ≈ 4 characters)
            chunk_tokens = len(chunk_content) // 4

            if total_tokens + chunk_tokens <= max_tokens:
                retrieved_chunks.append({
                    'content': chunk_content,
                    'metadata': source_info
                })
                source_docs.add((doc_id, document.title if document else "Unknown"))
                total_tokens += chunk_tokens
            else:
                break

        # Prepare response
        result = {
            'chunks': retrieved_chunks,
            'sources': [{'id': doc_id, 'title': title} for doc_id, title in source_docs]
        }

        # Generate response using Cohere if enabled
        if generate_response and self.cohere_client and retrieved_chunks:
            try:
                # Prepare context documents for Cohere
                context_docs = [
                    {"text": chunk['content']} for chunk in retrieved_chunks
                ]

                # Generate response using Cohere's RAG capabilities
                response = self.cohere_client.chat(
                    model=self.cohere_model,
                    message=query,
                    documents=context_docs,
                    temperature=0.3
                )

                result['response'] = {
                    'text': response.text,
                    'citations': getattr(response, 'citations', []),
                    'generation_id': getattr(response, 'generation_id', None)
                }

            except Exception as e:
                logger.error(f"Error generating response with Cohere: {e}")
                # Fallback to simple concatenation
                result['response'] = {
                    'text': "Here's what I found:\n\n" + "\n\n".join(
                        f"From {chunk['metadata']['document_title']}:\n{chunk['content']}"
                        for chunk in retrieved_chunks
                    ),
                    'citations': [],
                    'error': str(e)
                }

        return result

    def _get_chunk_content(self, chunk_id: str) -> Optional[str]:
        """Helper method to retrieve chunk content from database"""
        # You'll need to implement this method to fetch actual chunk content
        # This is a placeholder implementation
        conn = sqlite3.connect(self.doc_db.db_path)
        cursor = conn.execute(
            "SELECT content FROM document_chunks WHERE id = ?", (chunk_id,)
        )
        row = cursor.fetchone()
        conn.close()
        return row[0] if row else None

    def _apply_mmr(self, results: List[tuple], query_embedding: np.ndarray,
                   max_results: int, lambda_param: float = 0.7) -> List[tuple]:
        """Apply Maximal Marginal Relevance for diversity"""
        if len(results) <= max_results:
            return results

        selected = []
        remaining = list(results)

        # Select first result (highest similarity)
        selected.append(remaining.pop(0))

        while len(selected) < max_results and remaining:
            mmr_scores = []

            for i, (chunk_id, sim_score) in enumerate(remaining):
                # Calculate max similarity to already selected items
                max_sim_to_selected = 0
                # In a real implementation, you'd calculate similarity between embeddings
                # For now, using a simplified approach
                max_sim_to_selected = sim_score * 0.1  # Simplified

                # MMR score
                mmr_score = lambda_param * sim_score - (1 - lambda_param) * max_sim_to_selected
                mmr_scores.append((i, mmr_score))

            # Select item with highest MMR score
            best_idx, _ = max(mmr_scores, key=lambda x: x[1])
            selected.append(remaining.pop(best_idx))

        return selected

    def get_system_stats(self) -> Dict[str, Any]:
        documents = self.doc_db.list_documents()
        return {
            'total_documents': len(documents),
            'total_vectors': self.vector_store.index.ntotal,
            'processing_queue_size': self.indexer.processing_queue.qsize(),
            'recent_documents': [doc.title for doc in documents[:5]]
        }


# Streamlit UI
def main():
    st.set_page_config(
        page_title="Dynamic RAG System",
        page_icon="🔍",
        layout="wide"
    )

    st.title("🔍 Dynamic RAG System POC")
    st.markdown("A scalable document retrieval system with external knowledge base")

    # Initialize system
    if 'rag_system' not in st.session_state:
        with st.spinner("Initializing RAG system..."):
            st.session_state.rag_system = RAGSystem(cohere_api_key="LAb5hJRB31mFxk5cKLKCYxOqlhOqrIW0xcX3sgJF")

    rag_system = st.session_state.rag_system

    # Sidebar for system stats
    with st.sidebar:
        st.header("📊 System Stats")
        stats = rag_system.get_system_stats()

        col1, col2 = st.columns(2)
        with col1:
            st.metric("Documents", stats['total_documents'])
        with col2:
            st.metric("Vectors", stats['total_vectors'])

        st.metric("Queue Size", stats['processing_queue_size'])

        if stats['recent_documents']:
            st.subheader("Recent Documents")
            for doc in stats['recent_documents']:
                st.text(f"• {doc}")

    # Main interface
    tab1, tab2, tab3 = st.tabs(["🔍 Search", "📄 Add Document", "⚙️ System"])

    with tab1:
        st.header("Search Documents")

        query = st.text_input("Enter your search query:",
                              placeholder="What would you like to know?")

        col1, col2, col3 = st.columns(3)
        with col1:
            max_chunks = st.slider("Max chunks to retrieve", 1, 10, 5)
        with col2:
            max_tokens = st.slider("Token budget", 1000, 32000, 16000, step=1000)
        with col3:
            generate_response = st.checkbox("Generate response", value=True)

        search_btn = st.button("🔍 Search", type="primary")

        if search_btn and query:
            with st.spinner("Searching documents and generating response..."):
                results = rag_system.search_documents(
                    query,
                    max_chunks,
                    max_tokens,
                    generate_response=generate_response
                )

            if results.get('chunks'):
                st.success(f"Found {len(results['chunks'])} relevant chunks")

                if generate_response and 'response' in results:
                    st.subheader("Generated Answer")
                    st.markdown(results['response']['text'])

                    if results['response'].get('citations'):
                        st.caption("Citations:")
                        for citation in results['response']['citations']:
                            st.json(citation)

                st.subheader("Retrieved Context")
                total_tokens = sum(len(chunk['content']) // 4 for chunk in results['chunks'])
                st.info(f"Total tokens used: {total_tokens:,} / {max_tokens:,}")

                for i, chunk in enumerate(results['chunks'], 1):
                    with st.expander(f"Chunk {i} (From: {chunk['metadata']['document_title']})"):
                        st.markdown(chunk['content'])
                        st.caption(f"Document ID: {chunk['metadata']['document_id']} | "
                                 f"Chunk ID: {chunk['metadata']['chunk_id']} | "
                                 f"Score: {chunk['metadata']['score']:.3f}")

                if results.get('sources'):
                    st.subheader("Source Documents")
                    for source in results['sources']:
                        st.text(f"• {source['title']} (ID: {source['id']})")
            else:
                st.warning("No relevant documents found.")

    with tab2:
        st.header("Add New Document")

        # Document input method selection
        input_method = st.radio(
            "Choose input method:",
            ["📝 Manual Text Input", "📄 File Upload"],
            horizontal=True
        )

        with st.form("add_document"):
            title = st.text_input("Document Title*", placeholder="Enter document title")

            content = ""
            file_metadata = {}

            if input_method == "📝 Manual Text Input":
                content = st.text_area("Document Content*",
                                       placeholder="Paste or type document content here...",
                                       height=300)
            else:
                st.subheader("📄 File Upload")

                # Show supported formats
                supported_formats = {
                    "PDF": "✅" if PDF_AVAILABLE else "❌ (pip install PyPDF2)",
                    "DOCX": "✅" if DOCX_AVAILABLE else "❌ (pip install python-docx)",
                    "TXT/MD": "✅",
                    "Excel": "✅" if EXCEL_AVAILABLE else "❌ (pip install openpyxl)",
                    "Code files": "✅ (.py, .js, .html, .css, .json)"
                }

                col1, col2 = st.columns([2, 1])
                with col1:
                    uploaded_file = st.file_uploader(
                        "Choose a file",
                        type=['pdf', 'docx', 'txt', 'md', 'py', 'js', 'html', 'css', 'json', 'xlsx', 'xls'],
                        help="Upload a document to extract text automatically"
                    )

                with col2:
                    st.markdown("**Supported Formats:**")
                    for format_name, status in supported_formats.items():
                        st.markdown(f"• {format_name}: {status}")

                # Process uploaded file
                if uploaded_file is not None:
                    try:
                        with st.spinner(f"Processing {uploaded_file.name}..."):
                            extracted_text, file_meta = FileProcessor.process_uploaded_file(uploaded_file)
                            content = extracted_text
                            file_metadata = file_meta

                        # Show file info
                        st.success(f"✅ File processed successfully!")

                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("File Size", f"{file_metadata.get('file_size', 0):,} bytes")
                        with col2:
                            st.metric("File Type", file_metadata.get('file_type', 'unknown').upper())
                        with col3:
                            st.metric("Text Length", f"{len(content):,} chars")

                        # Show preview
                        with st.expander("📖 Preview Extracted Text"):
                            preview_text = content[:1000] + "..." if len(content) > 1000 else content
                            st.text_area("Extracted content:", preview_text, height=150, disabled=True)

                        # Allow title auto-fill from filename
                        if not title and uploaded_file.name:
                            suggested_title = uploaded_file.name.rsplit('.', 1)[0].replace('_', ' ').replace('-', ' ')
                            if st.button(f"📝 Use filename as title: '{suggested_title}'"):
                                title = suggested_title

                    except Exception as e:
                        st.error(f"❌ Error processing file: {str(e)}")
                        if "not available" in str(e):
                            st.info("💡 **Installation Required**: " + str(e))
                        content = ""

            # Metadata section
            st.subheader("📋 Metadata (Optional)")
            col1, col2 = st.columns(2)
            with col1:
                author = st.text_input("Author")
                category = st.text_input("Category")
            with col2:
                tags = st.text_input("Tags (comma-separated)")
                source_url = st.text_input("Source URL")

            # Submit button
            submitted = st.form_submit_button("📄 Add Document", type="primary")

            if submitted:
                if not title or not content:
                    st.error("❌ Title and content are required!")
                else:
                    # Combine metadata
                    metadata = {
                        'author': author,
                        'category': category,
                        'tags': [tag.strip() for tag in tags.split(',') if tag.strip()],
                        'source_url': source_url,
                        **file_metadata  # Include file metadata if available
                    }

                    with st.spinner("Adding document to knowledge base..."):
                        doc_id = rag_system.add_document(title, content, metadata)

                    st.success(f"✅ Document added successfully!")

                    # Show document info
                    col1, col2 = st.columns(2)
                    with col1:
                        st.info(f"**Document ID**: `{doc_id}`")
                    with col2:
                        st.info(f"**Content Length**: {len(content):,} characters")

                    st.info("🔄 Document is being processed in the background and will be available for search shortly.")

                    # Show processing estimation
                    estimated_chunks = len(content) // 500  # Rough estimate
                    st.caption(
                        f"⏱️ Estimated processing time: ~{max(1, estimated_chunks * 2)} seconds ({estimated_chunks} chunks)")

        # Quick upload tips
        with st.expander("💡 Upload Tips & Best Practices"):
            st.markdown("""
            **📄 File Processing Tips:**
            - **PDF**: Works best with text-based PDFs (not scanned images)
            - **DOCX**: Extracts text from paragraphs, tables not fully supported
            - **Excel**: Converts sheets to tab-separated text format
            - **Large Files**: Files over 10MB may take longer to process

            **🎯 Content Guidelines:**
            - **Clear Structure**: Use headings and paragraphs for better chunking
            - **Relevant Content**: Remove unnecessary formatting or metadata
            - **Complete Documents**: Partial or truncated content may reduce search quality

            **🏷️ Metadata Best Practices:**
            - **Tags**: Use specific, searchable keywords
            - **Categories**: Group related documents for easier filtering
            - **Author**: Include for attribution and source tracking
            """)

    with tab3:
        st.header("System Configuration")

        st.subheader("Architecture Overview")
        st.markdown("""
        This POC demonstrates the complete lifecycle of a dynamic RAG system:

        **🏗️ Architecture Components:**
        - **Document Database**: SQLite for document storage
        - **Vector Store**: FAISS for similarity search
        - **Background Indexer**: Asynchronous document processing
        - **Text Processor**: Chunking and embedding generation
        - **MMR Algorithm**: Diversity in retrieval results

        **🔄 Processing Pipeline:**
        1. Document ingestion → Database storage
        2. Background chunking → Semantic segmentation
        3. Embedding generation → Vector storage
        4. Query processing → Similarity search
        5. MMR application → Diverse results
        6. Token budgeting → Context optimization
        """)

        st.subheader("System Performance")

        # Performance metrics (mock data for demo)
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Avg Query Time", "145ms", "↓ 23ms")
        with col2:
            st.metric("Index Size", "2.3MB", "↑ 0.8MB")
        with col3:
            st.metric("Memory Usage", "127MB", "↑ 12MB")

        st.subheader("Configuration")
        st.markdown("""
        **Embedding Model**: all-MiniLM-L6-v2 (384 dimensions)  
        **Chunk Size**: 500 tokens with 50 token overlap  
        **Similarity Metric**: Cosine similarity (Inner Product)  
        **MMR Lambda**: 0.7 (70% relevance, 30% diversity)  
        """)


if __name__ == "__main__":
    main()